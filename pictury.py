from kivy.config import Config
Config.set('kivy','keyboard_mode','system')

Config.set('graphics','resizable', True)
Config.set('graphics','height', 440)
Config.set('graphics','resizable', True)
Config.set('graphics', 'width', 390)

from kivy.app import App
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
import re
import docx
from kivy.uix.gridlayout import GridLayout
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.anchorlayout import AnchorLayout
from kivy.properties import ObjectProperty
from kivy.uix.label import Label
from kivy.core.window import Window
from kivy.uix.spinner import Spinner,SpinnerOption


Window.clearcolor=(1,1,1,1)

App.title='Search in file'
class SearchApp(App):
        def build(self):
            self.lebel=ObjectProperty(None)
            self.text_input=ObjectProperty(None)
            self.button=ObjectProperty(None)
            self.spinner=SpinnerOption()
            self.box_layout=BoxLayout(orientation='vertical')
            self.anchor_layout=AnchorLayout(anchor_x='center',anchor_y='center',size_hint=(1,0.3))
            self.spinner=Spinner(background_normal='#3366FF',background_color='#3366FF',values=['txt','doc','docx'],text='Расширение текстового файла',size_hint=(0.1,0.1),color='white',font_size=21)
            self.lebel=Label(text='',size_hint=(1,0.01),font_size=30,strip=True,color='yellow',italic=True,outline_color=[57,189,92],outline_width=1,)
            self.grid_layout=GridLayout(rows=4,padding=[40,5,40,34],spacing=[5,20],size_hint=(1,0.7))
            self.text_input = TextInput(opacity=0.8,on_touch_down=self.on_down,size_hint=(0.5,0.1),multiline=False,font_size=22,background_color='#00FFFF',hint_text='полный путь к файлу поиска',foreground_color='0033FF')
            self.text_input1 = TextInput(opacity=0.8,on_touch_down=self.on_down,size_hint=(0.5,0.1), multiline=False,font_size=27,background_color='#00FFFF',hint_text='искомое слово')
            self.button=Button(text='поиск',size_hint=(0.5,0.45),color='#0033FF',background_color='#00FF99',background_normal='',on_press=self.fork,font_size=24,opacity=0.8)
            self.grid_layout.add_widget(self.text_input1)
            self.grid_layout.add_widget(self.text_input)
            self.grid_layout.add_widget(self.spinner)
            self.grid_layout.add_widget(self.lebel)
            self.anchor_layout.add_widget(self.button)
            self.box_layout.add_widget(self.grid_layout)
            self.box_layout.add_widget(self.anchor_layout)
            return self.box_layout

        def on_down(self,text,instance):
            self.lebel.text=''

        def fork(self,instance):
            try:
                text_2=str(self.text_input1.text)
                text_1=str(self.text_input.text)
                fakl=str(self.spinner.text)
                if fakl=='docx':
                    file=docx.Document(text_1)
                    for rand in file.paragraphs:
                       goo=str(rand.text)
                       moll=bool(re.search(text_2,goo))
                       if moll==True:
                           self.lebel.text='есть это слово'
                           break
                       else:
                           self.lebel.text = 'нет этого слово'
                elif fakl=='txt':
                    with open(text_1,'r',encoding='utf-8') as f:
                            strok = f.read()
                    result_ = bool(re.search(text_2, strok))
                    if result_ == True:
                        self.lebel.text = 'есть это слово'
                    else:
                        self.lebel.text = 'нет этого слова'
                else:
                    file = docx.Document(text_1)
                    for rand in file.paragraphs:
                        goo = str(rand.text)
                        moll = bool(re.search(text_2, goo))
                        if moll == True:
                            self.lebel.text = ''
                            break
            except Exception:
                self.lebel.text='!'

if __name__ == '__main__':
    SearchApp().run()
